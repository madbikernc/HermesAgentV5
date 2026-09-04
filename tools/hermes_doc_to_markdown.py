#!/usr/bin/env python3
# Version: 1.0.0
#
# 1.0.0 (2026-09-04) — direct request: detect whether a targeted document is
# PDF or DOCX and convert it to Markdown for RAG indexing/retrieval, rather
# than the flat plain-text extraction hermes-rag-ingest-kb.py used before
# (pypdf/python-docx page/paragraph joins with no structure at all). PDF
# uses pymupdf4llm (added to infra/hermes-rag/requirements.txt) — the
# standard library for this, doing font-size-based heading detection, table
# reconstruction, and list formatting; verified live against a real PDF
# already in RAGDocs (a GLOCK manual) before trusting it, including that its
# OCR is real and automatic. DOCX has no equivalent off-the-shelf converter
# worth adding a dependency for -- python-docx already exposes everything
# needed (paragraph styles for headings/lists, run-level bold/italic,
# tables), so docx_to_markdown() below is hand-rolled directly against it.
#
# use_ocr defaults to False here, matching hermes-rag-ingest-kb.py's own
# S16c precedent (OCR is real CPU/GPU cost — an operator's explicit --ocr
# choice per run, not a silent default) -- confirmed live that pymupdf4llm
# 1.28.2's own default is use_ocr=True (it OCR'd 8 pages unprompted on the
# first test run against a real scanned PDF before this was noticed and
# pinned down to a real kwarg, not assumed).
"""
hermes_doc_to_markdown.py — Detect whether a document is PDF or DOCX and
convert it to Markdown. Named with an underscore, breaking this project's
usual hyphenated tools/ filename convention, deliberately: this module is
both `import`ed (hermes-rag-ingest-kb.py, for every PDF/DOCX its recursive
RAGDocs scan finds) and runnable standalone on one file at a time — Python
cannot import a module whose filename contains a hyphen, same reasoning
hermes_rag_common.py's own header already documents.

Standalone usage:
    /opt/hermes/venvs/rag/bin/python3 hermes_doc_to_markdown.py DOC [--ocr] [-o OUT.md]
    (default output: DOC's own path with its suffix replaced by .md;
    "-o -" writes to stdout instead)

Library usage:
    from hermes_doc_to_markdown import to_markdown, detect_doc_type
    md_text = to_markdown(path, ocr=False)   # raises ValueError for an
                                              # unhandled extension
"""
import argparse
import sys
from pathlib import Path

DOC_EXTENSIONS = {".pdf": "pdf", ".docx": "docx"}


def detect_doc_type(path: Path) -> str | None:
    """Returns "pdf", "docx", or None for anything else. Extension-based,
    not content-sniffed -- every real caller here (the recursive RAGDocs
    scan, this module's own CLI) already knows the file's real type from
    how it found the path in the first place; sniffing magic bytes would
    only matter for a mislabeled/renamed file, which has never come up
    against this project's real content."""
    return DOC_EXTENSIONS.get(Path(path).suffix.lower())


def pdf_to_markdown(path: Path, ocr: bool = False) -> str:
    """pymupdf4llm does font-size-based heading detection, table
    reconstruction (GitHub-flavored markdown tables), and list formatting
    natively -- verified live against a real scanned/mixed PDF (a GLOCK
    manual in RAGDocs/Firearms) before trusting it, OCR included when
    ocr=True. ocr=False (the default) still gets every page with a real
    text layer; only image-only pages come back thin/empty, same tradeoff
    hermes-rag-ingest-kb.py's own pre-existing --ocr flag already made for
    its pypdf-based path."""
    import pymupdf4llm

    return pymupdf4llm.to_markdown(str(path), use_ocr=ocr).strip()


# ---- DOCX -> Markdown ------------------------------------------------------

def _iter_block_items(document):
    """Walk a python-docx Document's body in real reading order, paragraphs
    and tables interleaved -- document.paragraphs/document.tables alone
    each only give one type, in isolation from where it actually sits
    relative to the other (confirmed live against python-docx's own docs;
    this is the documented idiom for body-order iteration, not a guess)."""
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _docx_run_text(run) -> str:
    """One run's text, wrapped in markdown emphasis markers per its own
    bold/italic flags. A run with no visible text (a bare formatting
    change, common in Word's real internal run-splitting) contributes
    nothing rather than an empty **/_ pair."""
    text = run.text
    if not text:
        return ""
    if run.bold and run.italic:
        return f"***{text}***"
    if run.bold:
        return f"**{text}**"
    if run.italic:
        return f"_{text}_"
    return text


def _docx_paragraph_text(paragraph) -> str:
    return "".join(_docx_run_text(r) for r in paragraph.runs).strip()


def _docx_table_markdown(table) -> str:
    """GitHub-flavored markdown table. A cell's own internal paragraph
    breaks become <br> (a raw newline would break the row out of the table
    syntax entirely); a literal "|" is escaped for the same reason. First
    row is always treated as the header -- real Word tables in practice
    almost always have one, and a header-less table still renders fine,
    just with a generic first row as the header visually."""
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            text = "<br>".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
            cells.append(text.replace("|", "\\|").replace("\n", " "))
        rows.append(cells)
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    lines = ["| " + " | ".join(rows[0]) + " |", "| " + " | ".join(["---"] * width) + " |"]
    for r in rows[1:]:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


def docx_to_markdown(path: Path) -> str:
    """Heading N -> "#"*N; Title/Subtitle -> "#"/"##" (Word has no numbered
    equivalent, so these two are the only style names mapped by name
    instead of a parsed digit); "List Bullet"/"List Bullet N" -> a "-" item
    indented by (N-1) two-space levels, same shape for "List Number"/
    "List Number N" but with "1." (markdown renderers auto-number
    regardless of the literal digit used, so there's no need to track each
    list's real running count); anything else -> plain paragraph text.
    Tables are rendered in their real document position via
    _iter_block_items(), not appended at the end. Images are NOT extracted
    -- out of scope for a text/markdown RAG pipeline with no asset store to
    put them in; noted here rather than silently unhandled."""
    from docx import Document

    doc = Document(str(path))
    blocks = []
    for item in _iter_block_items(doc):
        cls_name = type(item).__name__
        if cls_name == "Table":
            table_md = _docx_table_markdown(item)
            if table_md:
                blocks.append(table_md)
            continue

        text = _docx_paragraph_text(item)
        if not text:
            continue
        style = item.style.name if item.style else "Normal"

        if style.startswith("Heading "):
            suffix = style[len("Heading "):].strip()
            level = int(suffix) if suffix.isdigit() else 1
            level = max(1, min(level, 6))
            blocks.append(f"{'#' * level} {text}")
        elif style == "Title":
            blocks.append(f"# {text}")
        elif style == "Subtitle":
            blocks.append(f"## {text}")
        elif style.startswith("List Bullet"):
            suffix = style[len("List Bullet"):].strip()
            depth = int(suffix) - 1 if suffix.isdigit() else 0
            blocks.append(f"{'  ' * max(depth, 0)}- {text}")
        elif style.startswith("List Number"):
            suffix = style[len("List Number"):].strip()
            depth = int(suffix) - 1 if suffix.isdigit() else 0
            blocks.append(f"{'  ' * max(depth, 0)}1. {text}")
        else:
            blocks.append(text)

    return "\n\n".join(blocks).strip()


# ---- dispatch ---------------------------------------------------------

def to_markdown(path: Path, ocr: bool = False) -> str:
    kind = detect_doc_type(path)
    if kind == "pdf":
        return pdf_to_markdown(path, ocr=ocr)
    if kind == "docx":
        return docx_to_markdown(path)
    raise ValueError(f"{path}: not a PDF or DOCX (unhandled extension {Path(path).suffix!r})")


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__.split("\n\n")[0])
    ap.add_argument("doc", metavar="DOC", help="Path to a .pdf or .docx file")
    ap.add_argument("-o", "--output", metavar="OUT",
                     help='Output .md path ("-" for stdout; default: DOC with .md suffix)')
    ap.add_argument("--ocr", action="store_true",
                     help="PDF only: OCR pages with no real text layer (off by default)")
    args = ap.parse_args()

    doc_path = Path(args.doc)
    if not doc_path.is_file():
        print(f"ERROR: {doc_path}: not a file", file=sys.stderr)
        return 1

    try:
        md = to_markdown(doc_path, ocr=args.ocr)
    except ValueError as e:
        print(f"ERROR: {e}", file=sys.stderr)
        return 1

    if not md.strip():
        print(f"WARNING: {doc_path}: no extractable content "
              f"(scanned/image-only PDF? pass --ocr)", file=sys.stderr)
        return 1

    if args.output == "-":
        print(md)
        return 0

    out_path = Path(args.output) if args.output else doc_path.with_suffix(".md")
    out_path.write_text(md, encoding="utf-8")
    print(f"{out_path}: {len(md)} chars")
    return 0


if __name__ == "__main__":
    sys.exit(main())
